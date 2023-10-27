import streamlit as st
import os
import docx

# Function to extract the first 250 characters from a text file
def extract_text_from_file(file):
    with open(file, 'r', encoding='utf-8') as f:
        content = f.read()
        return content[:250]

# Function to generate a Word document from extracted text
def generate_word_document(texts):
    doc = docx.Document()
    for i, text in enumerate(texts):
        doc.add_paragraph(f"Text from File {i + 1}:")
        doc.add_paragraph(text)
    return doc

# Streamlit interface
st.title("Text File Extractor and Word Document Generator")

uploaded_files = st.file_uploader("Upload multiple .txt files", type=["txt"], accept_multiple_files=True)

if uploaded_files:
    st.write("Uploaded files:")
    text_contents = []
    for uploaded_file in uploaded_files:
        text = extract_text_from_file(uploaded_file)
        text_contents.append(text)

    if st.button("Generate Word Document"):
        doc = generate_word_document(text_contents)
        doc.save("generated_document.docx")
        st.success("Word document generated successfully! You can download it below.")

        # Download link for the generated Word document
        with open("generated_document.docx", "rb") as docx_file:
            st.download_button(
                label="Download Word Document",
                data=docx_file,
                key="word_doc",
                file_name="generated_document.docx",
            )
else:
    st.info("Please upload some .txt files.")

# Cleanup: remove the generated Word document file
if os.path.exists("generated_document.docx"):
    os.remove("generated_document.docx")
